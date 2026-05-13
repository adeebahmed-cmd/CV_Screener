import io
import logging
import re
from pathlib import Path
from typing import Optional

import pdfplumber
from docx import Document

from config import MAX_FILE_MB

log = logging.getLogger(__name__)

MAX_BYTES = MAX_FILE_MB * 1024 * 1024

_HEADER_WORDS = {
    "curriculum", "curriculum vitae", "resume", "cv", "profile",
    "personal details", "personal information", "bio-data", "biodata",
}

_PROFILE_HINTS = (
    "summary", "profile", "experience", "employment", "work history", "skills",
    "competencies", "expertise", "responsibilities", "achievements", "education",
    "certification", "certifications", "projects", "leadership", "tools", "languages",
)

_PROFILE_KEYWORDS = (
    "experience", "years", "skill", "skills", "managed", "management", "lead", "led",
    "director", "program", "project", "therapy", "rehabilitation", "clinical", "health",
    "operations", "budget", "training", "compliance", "evaluation", "partnership",
    "certified", "certification", "master", "bachelor", "phd", "mba",
)


class ParseError(ValueError):
    pass


def _clean(text: str) -> str:
    text = text.replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _ocr_pdf(data: bytes) -> str:
    """OCR fallback for scanned/image-based PDFs."""
    try:
        from pdf2image import convert_from_bytes
        import pytesseract
    except ImportError:
        raise ParseError("OCR libraries not installed. Install pdf2image and pytesseract.")
    images = convert_from_bytes(data, dpi=200)
    log.info("parse_pdf: OCR got %d page image(s)", len(images))
    pages = []
    for i, img in enumerate(images):
        raw = pytesseract.image_to_string(img)
        log.info("parse_pdf: OCR page %d chars=%d preview=%r", i+1, len(raw), raw[:80])
        pages.append(raw)
    return _clean("\n\n".join(pages))


def parse_pdf(data: bytes) -> str:
    text = ""
    try:
        pages = []
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page in pdf.pages:
                pages.append(page.extract_text() or "")
        text = _clean("\n\n".join(pages))
    except Exception as e:
        log.warning("parse_pdf: pdfplumber failed (%s), trying OCR", e)

    if not text:
        log.info("parse_pdf: no text extracted, falling back to OCR")
        try:
            text = _ocr_pdf(data)
        except Exception as e:
            raise ParseError(f"Could not extract text from PDF: {e}")

    return text


def parse_docx(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return _clean("\n".join(parts))


def parse_txt(data: bytes) -> str:
    try:
        return _clean(data.decode("utf-8"))
    except UnicodeDecodeError:
        return _clean(data.decode("latin-1", errors="ignore"))


def extract_text(filename: str, data: bytes) -> str:
    if len(data) > MAX_BYTES:
        raise ParseError(f"File exceeds {MAX_FILE_MB} MB limit.")
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return parse_pdf(data)
    if ext == ".docx":
        return parse_docx(data)
    if ext == ".txt":
        return parse_txt(data)
    raise ParseError(f"Unsupported file type: {ext}. Use PDF, DOCX, or TXT.")


_DEGREE_ABBREVS = re.compile(
    r"\b(MBBS|BAMS|BHMS|BDS|MD|MS|MBA|MHA|MPH|MPhil|PhD|BSc|MSc|BE|BTech|MTech|"
    r"BCA|MCA|BBA|MCom|BCom|LLB|LLM|PGDM|PGDBA|DM|MCh|DNB|FRCS|MRCP)\b",
    re.IGNORECASE,
)

_NAME_RE = re.compile(
    r"^[A-Z][a-z]+([\s\-\.][A-Z][a-z]+){0,4}$"  # Title Case words, 1–5 parts
)
_SENTENCE_STOPWORDS = {
    "ensuring", "focused", "experienced", "skilled", "dedicated", "results",
    "driven", "motivated", "proven", "dynamic", "passionate", "strategic",
    "seeking", "looking", "responsible", "managing", "leading", "developing",
    "working", "providing", "supporting", "delivering", "building",
}


def _filename_to_name(filename: str) -> str:
    """Convert a filename like 'John_Smith_CV.pdf' → 'John Smith'."""
    stem = Path(filename).stem
    # Remove common suffixes: CV, Resume, Updated, Final, 2024, etc.
    stem = re.sub(r"(?i)[_\-\s]*(cv|resume|updated|final|new|\d{4}).*$", "", stem)
    name = re.sub(r"[_\-]+", " ", stem).strip()
    # Title-case if all caps or all lower
    if name == name.upper() or name == name.lower():
        name = name.title()
    return name or stem


def guess_candidate_name(text: str, fallback_filename: Optional[str] = None) -> str:
    # Always prefer filename — it's the most reliable source
    if fallback_filename and Path(fallback_filename).stem.strip():
        fname = _filename_to_name(fallback_filename)
        if fname and fname.lower() not in {"unknown", "cv", "resume"}:
            return fname

    # Fall back to scanning first 30 lines of CV text for a name-like line
    for line in text.splitlines()[:30]:
        stripped = line.strip()
        if not stripped or len(stripped) > 50:
            continue
        low = stripped.lower()
        # Skip header words, contact info, sentences
        if any(low.startswith(h) or low == h for h in _HEADER_WORDS):
            continue
        if "@" in stripped or "http" in low or "linkedin" in low:
            continue
        if "|" in stripped or "/" in stripped or ":" in stripped:
            continue
        if re.search(r"\d", stripped):
            continue
        if _DEGREE_ABBREVS.search(stripped):
            continue
        # Skip lines that start with a sentence stopword (summaries/objectives)
        first_word = low.split()[0] if low.split() else ""
        if first_word in _SENTENCE_STOPWORDS:
            continue
        # Must match Title Case name pattern
        if _NAME_RE.match(stripped):
            return stripped

    return "Unknown Candidate"


def build_candidate_profile(text: str, max_chars: int = 5500) -> str:
    lines = [line.strip(" -\t") for line in text.splitlines()]
    lines = [line for line in lines if line]

    selected: list[str] = []
    seen: set[str] = set()

    def add(line: str) -> None:
        low = line.lower()
        if low in seen:
            return
        if "@" in line or "http" in low:
            return
        if re.search(r"\+?\d[\d\s().-]{7,}\d", line):
            return
        seen.add(low)
        selected.append(line)

    for line in lines[:12]:
        if len(line) <= 160:
            add(line)

    for idx, line in enumerate(lines):
        low = line.lower()
        if any(hint in low for hint in _PROFILE_HINTS):
            add(line)
            for follow in lines[idx + 1: idx + 5]:
                if len(follow) <= 180:
                    add(follow)

    for line in lines:
        low = line.lower()
        if any(keyword in low for keyword in _PROFILE_KEYWORDS):
            add(line)

    profile = "\n".join(selected)
    if len(profile) <= max_chars:
        return profile

    trimmed: list[str] = []
    total = 0
    for line in selected:
        extra = len(line) + (1 if trimmed else 0)
        if total + extra > max_chars:
            break
        trimmed.append(line)
        total += extra
    return "\n".join(trimmed)


# ---------------------------------------------------------------------------
# Domain dictionary — terms to scan for in JD text, grouped by category.
# Each entry: (keyword, category, default_weight, default_type)
# Weight is conservative (5–7); LLM-extracted weights override these.
# ---------------------------------------------------------------------------
_DOMAIN_DICT: list[tuple[str, str, int, str]] = [
    # ── Domain expertise ────────────────────────────────────────────────────
    ("WASH", "Domain", 7, "must-have"),
    ("M&E", "Domain", 7, "must-have"),
    ("monitoring and evaluation", "Domain", 7, "must-have"),
    ("public health", "Domain", 7, "must-have"),
    ("community health", "Domain", 6, "good-to-have"),
    ("global health", "Domain", 6, "good-to-have"),
    ("clinical care", "Domain", 7, "must-have"),
    ("patient management", "Domain", 6, "must-have"),
    ("disability inclusion", "Domain", 6, "must-have"),
    ("rehabilitation", "Domain", 6, "must-have"),
    ("physiotherapy", "Domain", 7, "must-have"),
    ("occupational therapy", "Domain", 7, "must-have"),
    ("mental health", "Domain", 6, "good-to-have"),
    ("nutrition", "Domain", 6, "good-to-have"),
    ("epidemiology", "Domain", 7, "must-have"),
    ("humanitarian", "Domain", 6, "good-to-have"),
    ("development sector", "Domain", 5, "good-to-have"),
    ("NGO sector", "Domain", 5, "good-to-have"),
    ("microfinance", "Domain", 7, "must-have"),
    ("banking", "Domain", 6, "must-have"),
    ("insurance", "Domain", 6, "must-have"),
    ("real estate", "Domain", 6, "must-have"),
    ("manufacturing", "Domain", 6, "must-have"),
    ("supply chain", "Domain", 6, "must-have"),
    ("logistics", "Domain", 6, "must-have"),
    ("retail", "Domain", 5, "good-to-have"),
    ("e-commerce", "Domain", 6, "good-to-have"),
    ("hospitality", "Domain", 6, "must-have"),
    ("education sector", "Domain", 6, "must-have"),
    ("legal services", "Domain", 6, "must-have"),
    ("media", "Domain", 5, "good-to-have"),
    ("agriculture", "Domain", 6, "must-have"),
    ("climate change", "Domain", 6, "good-to-have"),
    ("environment", "Domain", 5, "good-to-have"),
    ("gender equality", "Domain", 6, "good-to-have"),
    ("child protection", "Domain", 7, "must-have"),
    ("safeguarding", "Domain", 7, "must-have"),

    # ── Capability ───────────────────────────────────────────────────────────
    ("project management", "Capability", 7, "must-have"),
    ("programme management", "Capability", 7, "must-have"),
    ("budget management", "Capability", 7, "must-have"),
    ("stakeholder engagement", "Capability", 6, "must-have"),
    ("proposal writing", "Capability", 6, "must-have"),
    ("report writing", "Capability", 5, "good-to-have"),
    ("fundraising", "Capability", 7, "must-have"),
    ("grant management", "Capability", 7, "must-have"),
    ("donor reporting", "Capability", 6, "must-have"),
    ("needs assessment", "Capability", 6, "good-to-have"),
    ("community mobilization", "Capability", 6, "good-to-have"),
    ("procurement", "Capability", 6, "must-have"),
    ("contract management", "Capability", 6, "must-have"),
    ("vendor management", "Capability", 5, "good-to-have"),
    ("advocacy", "Capability", 6, "good-to-have"),
    ("policy development", "Capability", 6, "good-to-have"),
    ("risk management", "Capability", 6, "must-have"),
    ("compliance", "Capability", 6, "must-have"),
    ("quality assurance", "Capability", 6, "must-have"),
    ("operations management", "Capability", 6, "must-have"),
    ("business development", "Capability", 6, "must-have"),
    ("sales management", "Capability", 6, "must-have"),
    ("marketing", "Capability", 5, "good-to-have"),
    ("digital marketing", "Capability", 6, "good-to-have"),
    ("content creation", "Capability", 5, "good-to-have"),
    ("client management", "Capability", 6, "must-have"),
    ("account management", "Capability", 6, "must-have"),
    ("partnership management", "Capability", 6, "must-have"),

    # ── Leadership ───────────────────────────────────────────────────────────
    ("team leadership", "Leadership", 7, "must-have"),
    ("people management", "Leadership", 7, "must-have"),
    ("strategic planning", "Leadership", 7, "must-have"),
    ("change management", "Leadership", 6, "good-to-have"),
    ("organizational development", "Leadership", 6, "good-to-have"),
    ("performance management", "Leadership", 6, "must-have"),
    ("staff supervision", "Leadership", 6, "must-have"),
    ("executive leadership", "Leadership", 7, "must-have"),
    ("board engagement", "Leadership", 6, "good-to-have"),
    ("cross-functional leadership", "Leadership", 6, "good-to-have"),

    # ── Finance & Governance ─────────────────────────────────────────────────
    ("financial management", "Finance & Governance", 7, "must-have"),
    ("financial reporting", "Finance & Governance", 6, "must-have"),
    ("budgeting", "Finance & Governance", 6, "must-have"),
    ("auditing", "Finance & Governance", 7, "must-have"),
    ("internal audit", "Finance & Governance", 7, "must-have"),
    ("financial planning", "Finance & Governance", 6, "must-have"),
    ("cost control", "Finance & Governance", 6, "must-have"),
    ("accounts payable", "Finance & Governance", 6, "must-have"),
    ("accounts receivable", "Finance & Governance", 6, "must-have"),
    ("tax compliance", "Finance & Governance", 6, "must-have"),
    ("treasury management", "Finance & Governance", 6, "must-have"),
    ("governance", "Finance & Governance", 6, "must-have"),
    ("donor compliance", "Finance & Governance", 7, "must-have"),
    ("resource mobilization", "Finance & Governance", 6, "must-have"),

    # ── Research & Innovation ────────────────────────────────────────────────
    ("data analysis", "Research & Innovation", 7, "must-have"),
    ("quantitative research", "Research & Innovation", 6, "good-to-have"),
    ("qualitative research", "Research & Innovation", 6, "good-to-have"),
    ("impact assessment", "Research & Innovation", 6, "good-to-have"),
    ("baseline study", "Research & Innovation", 5, "good-to-have"),
    ("data collection", "Research & Innovation", 5, "good-to-have"),
    ("statistical analysis", "Research & Innovation", 6, "must-have"),
    ("research design", "Research & Innovation", 6, "good-to-have"),
    ("literature review", "Research & Innovation", 5, "good-to-have"),
    ("knowledge management", "Research & Innovation", 5, "good-to-have"),
    ("innovation", "Research & Innovation", 5, "good-to-have"),
    ("evidence-based practice", "Research & Innovation", 6, "good-to-have"),

    # ── Learning & Development ───────────────────────────────────────────────
    ("training", "Learning & Development", 6, "good-to-have"),
    ("capacity building", "Learning & Development", 6, "good-to-have"),
    ("facilitation", "Learning & Development", 5, "good-to-have"),
    ("mentoring", "Learning & Development", 5, "good-to-have"),
    ("coaching", "Learning & Development", 5, "good-to-have"),
    ("curriculum development", "Learning & Development", 6, "good-to-have"),
    ("instructional design", "Learning & Development", 6, "good-to-have"),
    ("e-learning", "Learning & Development", 5, "good-to-have"),
    ("staff development", "Learning & Development", 5, "good-to-have"),

    # ── Tools / Systems ──────────────────────────────────────────────────────
    ("SPSS", "Tools / Systems", 6, "must-have"),
    ("Stata", "Tools / Systems", 6, "must-have"),
    ("R programming", "Tools / Systems", 6, "must-have"),
    ("Python", "Tools / Systems", 6, "must-have"),
    ("Power BI", "Tools / Systems", 6, "must-have"),
    ("Tableau", "Tools / Systems", 6, "must-have"),
    ("Excel", "Tools / Systems", 5, "good-to-have"),
    ("SAP", "Tools / Systems", 6, "must-have"),
    ("Oracle", "Tools / Systems", 6, "must-have"),
    ("Salesforce", "Tools / Systems", 6, "must-have"),
    ("QuickBooks", "Tools / Systems", 6, "must-have"),
    ("Tally", "Tools / Systems", 6, "must-have"),
    ("GIS", "Tools / Systems", 6, "must-have"),
    ("DHIS2", "Tools / Systems", 7, "must-have"),
    ("KOBO", "Tools / Systems", 6, "must-have"),
    ("ODK", "Tools / Systems", 6, "must-have"),
    ("HMIS", "Tools / Systems", 6, "must-have"),
    ("ERP", "Tools / Systems", 6, "must-have"),
    ("SQL", "Tools / Systems", 6, "must-have"),
    ("Navision", "Tools / Systems", 6, "must-have"),
    ("Dynamics 365", "Tools / Systems", 6, "must-have"),
    ("HubSpot", "Tools / Systems", 5, "good-to-have"),
    ("Jira", "Tools / Systems", 5, "good-to-have"),
    ("AutoCAD", "Tools / Systems", 6, "must-have"),

    # ── Qualifications / Certifications ──────────────────────────────────────
    ("MBA", "Domain", 6, "good-to-have"),
    ("PhD", "Domain", 6, "good-to-have"),
    ("MBBS", "Domain", 7, "must-have"),
    ("MD", "Domain", 7, "must-have"),
    ("BAMS", "Domain", 6, "must-have"),
    ("MPH", "Domain", 6, "good-to-have"),
    ("MPhil", "Domain", 5, "good-to-have"),
    ("CPA", "Finance & Governance", 7, "must-have"),
    ("CA", "Finance & Governance", 7, "must-have"),
    ("ACCA", "Finance & Governance", 7, "must-have"),
    ("CFA", "Finance & Governance", 7, "must-have"),
    ("PMP", "Capability", 6, "good-to-have"),
    ("PRINCE2", "Capability", 6, "good-to-have"),
    ("CIPS", "Capability", 6, "good-to-have"),
    ("Six Sigma", "Capability", 6, "good-to-have"),

    # ── Soft Skills (only when explicitly required) ───────────────────────────
    ("negotiation", "Soft Skills", 5, "good-to-have"),
    ("conflict resolution", "Soft Skills", 5, "good-to-have"),
    ("critical thinking", "Soft Skills", 4, "good-to-have"),
    ("problem solving", "Soft Skills", 4, "good-to-have"),
    ("interpersonal skills", "Soft Skills", 4, "good-to-have"),
    ("presentation skills", "Soft Skills", 4, "good-to-have"),
    ("written communication", "Soft Skills", 4, "good-to-have"),
    ("time management", "Soft Skills", 4, "good-to-have"),
]


def augment_keywords(jd_text: str, existing_keywords: list[dict]) -> list[dict]:
    """Scan JD text against the domain dictionary and return keywords the LLM missed.

    Only adds a keyword if:
      1. It appears (case-insensitive) in the JD text.
      2. It is not already covered by the LLM's extracted keyword list
         (checked by case-insensitive exact match or substring containment).
    """
    text_lower = jd_text.lower()

    # Build a set of already-known keywords (lower) for dedup
    existing_lower: set[str] = set()
    for kw in existing_keywords:
        word = (kw.get("keyword") or "").strip().lower()
        if word:
            existing_lower.add(word)

    added: list[dict] = []
    seen_added: set[str] = set()

    for term, category, weight, kw_type in _DOMAIN_DICT:
        term_lower = term.lower()
        # Skip if already in LLM list or already being added
        if term_lower in seen_added:
            continue
        already_covered = any(
            term_lower in ex or ex in term_lower
            for ex in existing_lower
        )
        if already_covered:
            continue
        # Only add if the term appears as a whole word/phrase in the JD (not as a substring).
        # e.g. "excel" must not match inside "excellence".
        if re.search(r'\b' + re.escape(term_lower) + r'\b', text_lower):
            added.append({
                "keyword": term,
                "category": category,
                "weight": weight,
                "type": "good-to-have",
            })
            seen_added.add(term_lower)

    return added
